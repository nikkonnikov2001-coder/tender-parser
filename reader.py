import os

from docx import Document
from pypdf import PdfReader

from tz_docs import is_tz_docx, is_tz_pdf


def _pdf_ocr_enabled() -> bool:
    return os.environ.get("PDF_OCR", "").strip().lower() in (
        "1",
        "true",
        "yes",
        "on",
    )


def _pdf_ocr_trigger_chars() -> int:
    """
    Если pypdf дал меньше символов — пробуем OCR (при PDF_OCR=1).
    0 = всегда вызывать OCR и брать более длинный из двух вариантов.
    """
    try:
        n = int(os.environ.get("PDF_OCR_TRIGGER_CHARS", "80"))
    except ValueError:
        n = 80
    return max(0, min(n, 50_000))


def extract_text_from_docx(file_path):
    doc = Document(file_path)
    text_content = []
    
    # 1. Читаем обычные абзацы
    for para in doc.paragraphs:
        if para.text.strip():
            text_content.append(para.text.strip())
            
    # 2. КРИТИЧЕСКИ ВАЖНО: Читаем таблицы (тут прячется всё ТЗ)
    for table in doc.tables:
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                clean_text = cell.text.strip().replace('\n', ' ')
                # Добавляем текст ячейки, если он не пустой и мы его еще не добавили
                # (в docx бывают объединенные ячейки, которые дублируют текст)
                if clean_text and clean_text not in row_data:
                    row_data.append(clean_text)
            if row_data:
                # Склеиваем строку таблицы через разделитель
                text_content.append(" | ".join(row_data))
                
    return "\n".join(text_content)


def _extract_text_from_pdf_pypdf(file_path: str) -> str:
    try:
        reader = PdfReader(file_path)
    except Exception:
        return ""
    parts: list[str] = []
    for page in reader.pages:
        try:
            t = page.extract_text()
        except Exception:
            continue
        if t and t.strip():
            parts.append(t.strip())
    return "\n".join(parts)


def extract_text_from_pdf(file_path: str) -> str:
    text = _extract_text_from_pdf_pypdf(file_path)
    if not _pdf_ocr_enabled():
        return text
    trigger = _pdf_ocr_trigger_chars()
    plain_len = len(text.strip())
    need_ocr = trigger == 0 or plain_len < trigger
    if not need_ocr:
        return text
    try:
        from pdf_ocr import extract_pdf_text_via_ocr
    except ImportError as e:
        print(f"⚠️ PDF_OCR включён, но не хватает зависимостей (pymupdf/pytesseract/Pillow): {e}")
        return text
    try:
        ocr_text = extract_pdf_text_via_ocr(file_path)
    except RuntimeError as e:
        print(f"⚠️ {e}")
        return text
    except Exception as e:
        print(f"⚠️ PDF OCR: сбой ({e})")
        return text
    ocr_len = len(ocr_text.strip())
    if ocr_len > plain_len:
        print(
            f"📷 PDF OCR: использовано {ocr_len} симв. "
            f"(pypdf дал {plain_len}, порог PDF_OCR_TRIGGER_CHARS={trigger})."
        )
        return ocr_text
    if ocr_text.strip():
        print(
            "📷 PDF OCR: результат не длиннее pypdf — оставляем извлечение без OCR "
            f"({plain_len} vs {ocr_len} симв.)."
        )
    elif trigger == 0 or plain_len < trigger:
        print("📷 PDF OCR: пустой результат, остаётся текст pypdf.")
    return text


def get_tz_text(base_dir="downloads"):
    print("🔍 Начинаем глубокий поиск и чтение ТЗ...\n")
    
    for tender_id in os.listdir(base_dir):
        tender_path = os.path.join(base_dir, tender_id)
        
        if not os.path.isdir(tender_path):
            continue
            
        print(f"📁 Тендер: {tender_id}")
        tz_found = False
        
        for file in os.listdir(tender_path):
            if not (is_tz_docx(file) or is_tz_pdf(file)):
                continue
            file_path = os.path.join(tender_path, file)
            print(f"   📄 Читаем файл: {file}")

            try:
                if is_tz_docx(file):
                    full_text = extract_text_from_docx(file_path)
                else:
                    full_text = extract_text_from_pdf(file_path)

                print(f"   ✅ Текст вытащен! Длина: {len(full_text)} символов.")
                print(f"   📝 Превью:\n   {full_text[:300]}...\n")
                tz_found = True

            except Exception as e:
                print(f"   ❌ Ошибка чтения: {e}\n")

        if not tz_found:
            print(
                "   ⚠️ Нет подходящего .docx/.pdf по имени (маркеры в tz_docs), "
                "или старый .doc / скан без текста\n"
            )

if __name__ == "__main__":
    get_tz_text()